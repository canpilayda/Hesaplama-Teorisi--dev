import fitz  # PyMuPDF
import json
import os
import re
from PIL import Image
import pytesseract
from docx import Document

# ====== AYARLAR ======
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# ---------------------


def metin_temizle(metin):
    """
    AI eğitimi ve XML (Word) çıktısı için metni temizler.
    NULL byte, kontrol karakterleri ve tireleme hatalarını düzeltir.
    """
    if not metin:
        return ""

    # 1. ADIM: XML'i patlatan "Gulyabani" karakterleri (NULL bytes, control chars) temizle
    # \x00-\x08 : Null byte ve başındakiler
    # \x0b\x0c : Dikey tab ve form feed
    # \x0e-\x1f : Diğer kontrol karakterleri
    # (\n, \t ve \r karakterlerine dokunmuyoruz çünkü onlar lazım)
    metin = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", metin)

    # 2. ADIM: Hatalı bölünen kelimeleri birleştirme (Satır sonu tireleri)
    metin = re.sub(r"(\w+)-\s*\n\s*(\w+)", r"\1\2", metin)

    # 3. ADIM: Genel satır sonlarını ve fazla boşlukları temizleme
    metin = metin.replace("\n", " ").replace("\r", "")
    metin = re.sub(r"\s+", " ", metin).strip()

    return metin


# ... (Importlar ve metin_temizle aynı kalacak) ...


def pdf_isleyici_tam(pdf_path, progress_callback=None):
    """
    progress_callback: İlerleme durumunu (% olarak) bildirecek fonksiyon.
    """
    temizlenmis_bloklar = []
    try:
        doc = fitz.open(pdf_path)
        toplam_sayfa = doc.page_count

        for page_num in range(toplam_sayfa):
            # --- YÜZDE HESABI VE BİLDİRİM ---
            if progress_callback:
                yuzde = int((page_num / toplam_sayfa) * 100)
                progress_callback(yuzde)
            # -------------------------------

            page = doc.load_page(page_num)
            metin = page.get_text("text")

            # OCR Kontrolü
            if len(metin.strip()) < 10:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                temp_img_path = f"temp_{os.getpid()}_{page_num}.png"
                try:
                    pix.save(temp_img_path)
                    metin = pytesseract.image_to_string(
                        Image.open(temp_img_path), lang="eng"
                    )
                finally:
                    if os.path.exists(temp_img_path):
                        os.remove(temp_img_path)

            temiz_metin = metin_temizle(metin)
            if temiz_metin:
                temizlenmis_bloklar.append(
                    {"page_num": page_num + 1, "text": temiz_metin}
                )

        doc.close()

        # İşlem bittiğinde %100 gönder
        if progress_callback:
            progress_callback(100)

        return temizlenmis_bloklar

    except Exception as e:
        print(f"Hata: {e}")
        return []


# ... (Diğer kayıt fonksiyonları aynı kalacak) ...


def json_olarak_kaydet(veri_listesi, kayit_yolu):
    with open(kayit_yolu, "w", encoding="utf-8") as f:
        json.dump(veri_listesi, f, ensure_ascii=False, indent=4)


def word_olarak_kaydet(veri_listesi, kayit_yolu):
    doc = Document()
    doc.add_heading("PDF Metin Çıktısı", 0)
    for item in veri_listesi:
        doc.add_heading(f"Sayfa {item['page_num']}", level=1)
        doc.add_paragraph(item["text"])
        doc.add_paragraph("-" * 20)
    doc.save(kayit_yolu)


# =======================================================
# SENİN İSTEDİĞİN "TEK NOKTADAN YÖNETİM" FONKSİYONU
# =======================================================
def cikti_al(pdf_path, cikti_tipi, hedef_klasor="downloads"):
    """
    Tüm süreci tek başına yönetir.

    Parametreler:
    - pdf_path: İşlenecek PDF dosyasının yolu.
    - cikti_tipi: 'json' veya 'word'.
    - hedef_klasor: Dosyanın kaydedileceği klasör (varsayılan: downloads).

    Dönüş:
    - Oluşturulan dosyanın tam yolu (string) veya hata varsa None.
    """

    # 1. Önce veriyi analiz et (OCR vs yap)
    analiz_verisi = pdf_isleyici_tam(pdf_path)

    if not analiz_verisi:
        print("Veri çıkarılamadı veya dosya boş.")
        return None

    # 2. Dosya ismini hazırla (uzantısız)
    dosya_adi = os.path.splitext(os.path.basename(pdf_path))[0]
    os.makedirs(hedef_klasor, exist_ok=True)  # Klasör yoksa oluştur

    # 3. İstenen tipe göre kaydet
    if cikti_tipi.lower() == "json":
        cikti_yolu = os.path.join(hedef_klasor, dosya_adi + ".json")
        json_olarak_kaydet(analiz_verisi, cikti_yolu)
        return cikti_yolu

    elif cikti_tipi.lower() == "word":
        cikti_yolu = os.path.join(hedef_klasor, dosya_adi + ".docx")
        word_olarak_kaydet(analiz_verisi, cikti_yolu)
        return cikti_yolu

    else:
        print("Hatalı çıktı tipi! Sadece 'json' veya 'word' kullanabilirsin.")
        return None
